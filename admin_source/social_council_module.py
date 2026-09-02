# -*- coding: utf-8 -*-
"""Independent, zone-based Social Council management module."""

import os
import tempfile
from datetime import datetime

from PyQt5.QtCore import Qt, pyqtSignal, QSize
from PyQt5.QtGui import QPainter, QColor, QFont, QPen, QPixmap
from PyQt5.QtWidgets import (
    QWidget, QDialog, QVBoxLayout, QHBoxLayout, QFormLayout, QGridLayout,
    QLabel, QPushButton, QComboBox, QLineEdit, QTextEdit, QSpinBox,
    QDoubleSpinBox, QTableWidget, QTableWidgetItem, QHeaderView,
    QAbstractItemView, QMessageBox, QTabWidget, QGroupBox, QScrollArea,
    QFileDialog, QFrame, QProgressBar, QCheckBox, QInputDialog, QApplication
)

from header_widget import build_official_header
from icon_manager import get_icon, set_button_style
from jalali_utils import to_persian_digits, convert_dates_in_text
from database_social_council import (
    SOCIAL_ISSUE_CATEGORIES, SOCIAL_TARGET_GROUPS, SOCIAL_CONFIDENTIALITY_LEVELS,
)
from social_chart_reports import (
    CHART_TYPES, CLOSED_STATUSES, ACTIVE_ACTION_STATUSES,
    filter_rows, issues_by_category_payload, blocks_comparison_payload,
    committees_performance_payload, resolutions_status_payload, actions_status_payload,
)


def _text(value):
    return "" if value is None else str(value)


def _item(value, data=None):
    it = QTableWidgetItem(convert_dates_in_text(_text(value)))
    if data is not None:
        it.setData(Qt.UserRole, data)
    return it


def _selected_id(table, column=0):
    row = table.currentRow()
    if row < 0:
        return None
    value = table.item(row, column).data(Qt.UserRole)
    return value


def _table(headers):
    t = QTableWidget(0, len(headers))
    t.setHorizontalHeaderLabels(headers)
    t.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    t.setSelectionBehavior(QAbstractItemView.SelectRows)
    t.setSelectionMode(QAbstractItemView.SingleSelection)
    t.setEditTriggers(QAbstractItemView.NoEditTriggers)
    t.setAlternatingRowColors(True)
    return t


def _scroll(widget):
    area = QScrollArea()
    area.setWidgetResizable(True)
    area.setFrameShape(QFrame.NoFrame)
    area.setWidget(widget)
    return area


def _dialog_form(dialog):
    """Creates a resize-safe, vertically scrollable form for data-entry dialogs."""
    root = QVBoxLayout(dialog)
    root.setContentsMargins(8, 8, 8, 8)
    content = QWidget()
    form = QFormLayout(content)
    form.setFieldGrowthPolicy(QFormLayout.AllNonFixedFieldsGrow)
    form.setRowWrapPolicy(QFormLayout.WrapLongRows)
    area = _scroll(content)
    area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
    area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
    root.addWidget(area, 1)
    dialog._form_scroll_area = area
    dialog._form_content = content
    return form


class BarChartWidget(QWidget):
    """Dependency-free grouped horizontal bar chart for Persian management reports."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.title = "گزارش نموداری"
        self.subtitle = ""
        self.categories = []
        self.series = []
        self.setMinimumSize(900, 520)

    def set_chart(self, title, categories, series, subtitle=""):
        self.title = title or "گزارش نموداری"
        self.subtitle = subtitle or ""
        self.categories = list(categories or [])
        self.series = list(series or [])
        per_row = max(46, 24 * max(1, len(self.series)) + 12)
        self.setMinimumHeight(max(520, 155 + len(self.categories) * per_row))
        self.updateGeometry()
        self.update()

    def sizeHint(self):
        return QSize(1180, max(560, self.minimumHeight()))

    def save_png(self, path):
        old_size = self.size()
        target = QSize(max(1280, old_size.width()), max(720, self.minimumHeight()))
        self.resize(target)
        QApplication.processEvents()
        pixmap = QPixmap(target)
        pixmap.fill(Qt.white)
        painter = QPainter(pixmap)
        self.render(painter)
        painter.end()
        ok = pixmap.save(path, "PNG")
        self.resize(old_size)
        QApplication.processEvents()
        return ok

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing, True)
        painter.fillRect(self.rect(), QColor("#ffffff"))
        width, height = self.width(), self.height()
        painter.setPen(QColor("#17365d"))
        title_font = QFont(self.font()); title_font.setPointSize(15); title_font.setBold(True)
        painter.setFont(title_font)
        painter.drawText(24, 18, width - 48, 36, Qt.AlignCenter, self.title)
        top = 62
        if self.subtitle:
            subtitle_font = QFont(self.font()); subtitle_font.setPointSize(9)
            painter.setFont(subtitle_font); painter.setPen(QColor("#52616b"))
            painter.drawText(30, 52, width - 60, 38, Qt.AlignCenter | Qt.TextWordWrap, self.subtitle)
            top = 96
        if not self.categories or not self.series:
            painter.setPen(QColor("#7b8794")); painter.drawText(self.rect(), Qt.AlignCenter, "برای فیلتر انتخاب‌شده داده‌ای وجود ندارد.")
            return
        colors = [QColor("#1f77b4"), QColor("#2ca02c"), QColor("#ff8c00"), QColor("#8a5cf6"), QColor("#d14b68")]
        legend_x = 36
        painter.setFont(QFont(self.font().family(), 9))
        for index, series in enumerate(self.series):
            color = colors[index % len(colors)]
            painter.fillRect(legend_x, top, 16, 10, color)
            painter.setPen(QColor("#263238"))
            label = str(series.get("label") or f"سری {index + 1}")
            painter.drawText(legend_x + 22, top - 5, 160, 22, Qt.AlignLeft | Qt.AlignVCenter, label)
            legend_x += 190
        top += 34
        bottom = 42
        label_width = min(330, max(210, width // 4))
        chart_left = 58
        chart_right = max(chart_left + 250, width - label_width - 28)
        chart_width = max(1, chart_right - chart_left)
        values = []
        for series in self.series:
            values.extend(float(v or 0) for v in series.get("values", []))
        max_value = max(values) if values else 0
        if max_value <= 0:
            max_value = 1
        available_height = max(120, height - top - bottom)
        row_height = max(44, available_height / max(1, len(self.categories)))
        series_count = max(1, len(self.series))
        bar_height = max(8, min(20, (row_height - 10) / series_count))
        grid_pen = QPen(QColor("#e0e6eb")); grid_pen.setWidth(1)
        axis_font = QFont(self.font().family(), 8)
        painter.setFont(axis_font)
        for tick in range(6):
            x = chart_left + chart_width * tick / 5
            painter.setPen(grid_pen); painter.drawLine(int(x), top, int(x), int(top + row_height * len(self.categories)))
            painter.setPen(QColor("#607080"))
            label = to_persian_digits(f"{max_value * tick / 5:.0f}")
            painter.drawText(int(x - 24), int(top + row_height * len(self.categories) + 6), 48, 22, Qt.AlignCenter, label)
        label_font = QFont(self.font().family(), 9); label_font.setBold(True)
        value_font = QFont(self.font().family(), 8)
        for row, category in enumerate(self.categories):
            y0 = top + row * row_height
            painter.setFont(label_font); painter.setPen(QColor("#263238"))
            painter.drawText(chart_right + 10, int(y0), label_width - 18, int(row_height), Qt.AlignRight | Qt.AlignVCenter | Qt.TextWordWrap, str(category))
            for index, series in enumerate(self.series):
                vals = list(series.get("values") or [])
                value = float(vals[row] if row < len(vals) else 0)
                y = y0 + 5 + index * bar_height
                bar_width = chart_width * value / max_value
                painter.fillRect(int(chart_left), int(y), max(1, int(bar_width)), max(5, int(bar_height - 3)), colors[index % len(colors)])
                painter.setFont(value_font); painter.setPen(QColor("#263238"))
                text = to_persian_digits(f"{value:g}")
                painter.drawText(int(min(chart_right - 48, chart_left + bar_width + 6)), int(y - 2), 54, int(bar_height + 3), Qt.AlignLeft | Qt.AlignVCenter, text)
        painter.setPen(QPen(QColor("#9aa8b3"), 1))
        painter.drawLine(chart_left, int(top + row_height * len(self.categories)), chart_right, int(top + row_height * len(self.categories)))


class MemberDialog(QDialog):
    def __init__(self, db, zone_id, item=None, parent=None):
        super().__init__(parent)
        self.db, self.zone_id, self.item = db, zone_id, item or {}
        self.setWindowTitle("عضو شورای اجتماعی")
        self.resize(620, 610)
        form = _dialog_form(self)
        self.full_name = QLineEdit(self.item.get("full_name") or "")
        self.national_code = QLineEdit(self.item.get("national_code") or "")
        self.mobile = QLineEdit(self.item.get("mobile") or "")
        self.role = QComboBox(); self.role.setEditable(True)
        self.role.addItems(["رئیس شورا", "دبیر شورا", "نماینده کمیته", "معتمد محله", "عضو شورای اجتماعی", "نماینده دستگاه", "کارشناس اجتماعی"])
        self.role.setCurrentText(self.item.get("role_title") or "عضو شورای اجتماعی")
        self.rep_type = QComboBox(); self.rep_type.setEditable(True)
        self.rep_type.addItems(["عضو مردمی", "معتمد/شورای محله", "نماینده کمیته", "نماینده دستگاه", "مسئول مکان", "مدعو"])
        self.rep_type.setCurrentText(self.item.get("representation_type") or "عضو مردمی")
        self.committee = QComboBox(); self.committee.addItem("— بدون کمیته —", None)
        for c in db.get_zone_committees(zone_id, ensure=True): self.committee.addItem(c["title"], c["id"])
        idx = self.committee.findData(self.item.get("committee_id")); self.committee.setCurrentIndex(max(0, idx))
        self.agency = QLineEdit(self.item.get("agency_name") or "")
        self.start_date = QLineEdit(self.item.get("start_date") or "")
        self.end_date = QLineEdit(self.item.get("end_date") or "")
        self.status = QComboBox(); self.status.addItems(["فعال", "غیرفعال", "پایان همکاری"]); self.status.setCurrentText(self.item.get("status") or "فعال")
        self.notes = QTextEdit(self.item.get("notes") or ""); self.notes.setMinimumHeight(90)
        for label, w in [
            ("نام و نام خانوادگی:", self.full_name), ("کد ملی:", self.national_code), ("تلفن همراه:", self.mobile),
            ("سمت در شورا:", self.role), ("نوع نمایندگی:", self.rep_type), ("کمیته مرتبط:", self.committee),
            ("دستگاه/نهاد:", self.agency), ("شروع عضویت:", self.start_date), ("پایان عضویت:", self.end_date),
            ("وضعیت:", self.status), ("توضیحات:", self.notes),
        ]: form.addRow(label, w)
        buttons = QHBoxLayout(); save=set_button_style(QPushButton("ذخیره"), "save", "primary"); cancel=QPushButton("انصراف")
        save.clicked.connect(self.accept); cancel.clicked.connect(self.reject); buttons.addStretch(); buttons.addWidget(cancel); buttons.addWidget(save); form.addRow(buttons)

    def data(self):
        return dict(full_name=self.full_name.text().strip(), national_code=self.national_code.text().strip(),
                    mobile=self.mobile.text().strip(), role_title=self.role.currentText().strip(),
                    representation_type=self.rep_type.currentText().strip(), committee_id=self.committee.currentData(),
                    agency_name=self.agency.text().strip(), start_date=self.start_date.text().strip() or None,
                    end_date=self.end_date.text().strip() or None, status=self.status.currentText(), notes=self.notes.toPlainText().strip())


class MeetingDialog(QDialog):
    def __init__(self, db, zone_id, item=None, parent=None):
        super().__init__(parent); self.db=db; self.zone_id=zone_id; self.item=item or {}
        self.setWindowTitle("جلسه شورای اجتماعی"); self.resize(760, 760); form=_dialog_form(self)
        self.title=QLineEdit(self.item.get("title") or "جلسه شورای اجتماعی")
        self.date=QLineEdit(self.item.get("meeting_date") or ""); self.time=QLineEdit(self.item.get("start_time") or "")
        self.place=QComboBox(); self.place.setEditable(True); self.place.addItem("— انتخاب مکان —", None)
        for p in db.get_places(zone_id=zone_id):
            self.place.addItem(f"{p.get('name')} — {p.get('category') or p.get('subtype') or 'مکان'}", ("place", p.get("id"), p.get("name")))
        for m in db.get_mosques(zone_id=zone_id):
            self.place.addItem(f"{m.get('name')} — مسجد", ("mosque", m.get("id"), m.get("name")))
        for f in db.get_schools(zone_id=zone_id):
            self.place.addItem(f"{f.get('name')} — مدرسه", ("school", f.get("id"), f.get("name")))
        for f in db.get_health_centers(zone_id=zone_id):
            self.place.addItem(f"{f.get('name')} — مرکز بهداشتی", ("health_center", f.get("id"), f.get("name")))
        target = (self.item.get("place_source") or "place", self.item.get("place_ref_id") or self.item.get("place_id"), self.item.get("place_name"))
        idx = next((i for i in range(self.place.count()) if self.place.itemData(i) and self.place.itemData(i)[:2] == target[:2]), -1)
        if idx>=0:self.place.setCurrentIndex(idx)
        elif self.item.get("place_name"):self.place.setCurrentText(self.item["place_name"])
        self.agenda=QTextEdit(self.item.get("agenda") or "")
        self.attendees=QTextEdit(self.item.get("attendees") or "")
        self.absentees=QTextEdit(self.item.get("absentees") or "")
        self.invitees=QTextEdit(self.item.get("invitees") or "")
        self.minutes=QTextEdit(self.item.get("minutes_text") or "")
        self.attachment=QLineEdit(self.item.get("attachment_path") or "")
        browse=QPushButton("انتخاب فایل"); browse.clicked.connect(self._browse); aw=QHBoxLayout(); aw.addWidget(self.attachment,1); aw.addWidget(browse)
        self.status=QComboBox(); self.status.addItems(["برنامه‌ریزی‌شده", "برگزارشده", "لغوشده"]); self.status.setCurrentText(self.item.get("status") or "برنامه‌ریزی‌شده")
        form.addRow("عنوان:",self.title); form.addRow("تاریخ:",self.date); form.addRow("ساعت:",self.time); form.addRow("محل جلسه:",self.place)
        form.addRow("دستور جلسه:",self.agenda); form.addRow("حاضرین:",self.attendees); form.addRow("غایبین:",self.absentees)
        form.addRow("مدعوین:",self.invitees); form.addRow("صورت‌جلسه:",self.minutes); form.addRow("پیوست:",aw); form.addRow("وضعیت:",self.status)
        buttons=QHBoxLayout(); save=set_button_style(QPushButton("ذخیره"),"save","primary"); cancel=QPushButton("انصراف")
        save.clicked.connect(self.accept); cancel.clicked.connect(self.reject); buttons.addStretch(); buttons.addWidget(cancel); buttons.addWidget(save); form.addRow(buttons)
    def _browse(self):
        path,_=QFileDialog.getOpenFileName(self,"انتخاب پیوست")
        if path:self.attachment.setText(path)
    def data(self):
        selected = self.place.currentData()
        source, ref_id, saved_name = (selected if isinstance(selected, tuple) else ("place", None, self.place.currentText().split(" — ")[0].strip()))
        return dict(title=self.title.text().strip(),meeting_date=self.date.text().strip() or None,start_time=self.time.text().strip(),
                    place_id=(ref_id if source == "place" else None),place_source=source,place_ref_id=ref_id,
                    place_name=saved_name or self.place.currentText().split(" — ")[0].strip(),agenda=self.agenda.toPlainText().strip(),
                    attendees=self.attendees.toPlainText().strip(),absentees=self.absentees.toPlainText().strip(),invitees=self.invitees.toPlainText().strip(),
                    minutes_text=self.minutes.toPlainText().strip(),attachment_path=self.attachment.text().strip(),status=self.status.currentText())


class IssueDialog(QDialog):
    def __init__(self, db, item=None, parent=None):
        super().__init__(parent); self.db=db; self.item=item or {}; self.setWindowTitle("پرونده مسئله اجتماعی"); self.resize(720,760); form=_dialog_form(self)
        self.title=QLineEdit(self.item.get("title") or "")
        self.category=QComboBox(); self.category.setEditable(False)
        self.category.addItems(self.db.get_social_issue_categories())
        saved_category = self.item.get("category") or "سایر"
        if self.category.findText(saved_category) < 0:
            self.category.addItem(saved_category)
        self.category.setCurrentText(saved_category)
        self.add_category_btn=QPushButton("+"); self.add_category_btn.setFixedWidth(42)
        self.add_category_btn.setToolTip("افزودن دسته‌بندی جدید آسیب‌های اجتماعی")
        self.add_category_btn.clicked.connect(self._add_category)
        category_box=QHBoxLayout(); category_box.setContentsMargins(0,0,0,0); category_box.addWidget(self.category,1); category_box.addWidget(self.add_category_btn)
        self.urgency=QComboBox(); self.urgency.addItems(["کم","عادی","زیاد","فوری","بحرانی"]); self.urgency.setCurrentText(self.item.get("urgency") or "عادی")
        self.target=QComboBox(); self.target.setEditable(True); self.target.addItems(SOCIAL_TARGET_GROUPS); self.target.setCurrentText(self.item.get("target_group") or "عموم ساکنان")
        self.people=QSpinBox(); self.people.setRange(0,10000000); self.people.setValue(int(self.item.get("affected_people") or 0))
        self.households=QSpinBox(); self.households.setRange(0,10000000); self.households.setValue(int(self.item.get("affected_households") or 0))
        self.description=QTextEdit(self.item.get("description") or ""); self.evidence=QTextEdit(self.item.get("evidence") or "")
        self.source=QLineEdit(self.item.get("source") or "ثبت شورای اجتماعی"); self.agency=QLineEdit(self.item.get("responsible_agency") or "")
        self.status=QComboBox(); self.status.addItems(["ثبت اولیه","در حال بررسی","ارجاع‌شده","در حال اقدام","مختومه","انجام‌شده"]); self.status.setCurrentText(self.item.get("status") or "ثبت اولیه")
        self.conf=QComboBox(); self.conf.addItems(SOCIAL_CONFIDENTIALITY_LEVELS); self.conf.setCurrentText(self.item.get("confidentiality") or "داخلی")
        self.location=QLineEdit(self.item.get("location_text") or ""); self.lat=QLineEdit(_text(self.item.get("lat"))); self.lon=QLineEdit(_text(self.item.get("lon"))); self.due=QLineEdit(self.item.get("due_date") or "")
        self.mirror=QCheckBox("هم‌زمان در پرونده مسائل عمومی محله نیز ثبت شود"); self.mirror.setChecked(not bool(self.item))
        for label,w in [("عنوان:",self.title),("دسته‌بندی آسیب‌های اجتماعی:",category_box),("فوریت:",self.urgency),("گروه هدف:",self.target),
                        ("افراد درگیر:",self.people),("خانوارهای درگیر:",self.households),("شرح:",self.description),
                        ("مستندات/شواهد:",self.evidence),("منبع:",self.source),("دستگاه مسئول:",self.agency),
                        ("وضعیت:",self.status),("محرمانگی:",self.conf),("نشانی:",self.location),("عرض جغرافیایی:",self.lat),
                        ("طول جغرافیایی:",self.lon),("مهلت پیگیری:",self.due)]:form.addRow(label,w)
        if not self.item:form.addRow(self.mirror)
        buttons=QHBoxLayout();save=set_button_style(QPushButton("ذخیره"),"save","primary");cancel=QPushButton("انصراف");save.clicked.connect(self.accept);cancel.clicked.connect(self.reject);buttons.addStretch();buttons.addWidget(cancel);buttons.addWidget(save);form.addRow(buttons)
    def _add_category(self):
        title, ok = QInputDialog.getText(self, "دسته‌بندی جدید", "عنوان دسته‌بندی آسیب اجتماعی:")
        title = " ".join((title or "").split()).strip()
        if not ok or not title:
            return
        try:
            self.db.add_social_issue_category(title)
            self.category.clear()
            self.category.addItems(self.db.get_social_issue_categories())
            self.category.setCurrentText(title)
        except Exception as exc:
            QMessageBox.warning(self, "خطا", str(exc))
    def data(self):
        def number(v):
            try:return float(v.strip()) if v.strip() else None
            except:return None
        return dict(title=self.title.text().strip(),category=self.category.currentText().strip(),urgency=self.urgency.currentText(),target_group=self.target.currentText().strip(),
                    affected_people=self.people.value(),affected_households=self.households.value(),description=self.description.toPlainText().strip(),
                    evidence=self.evidence.toPlainText().strip(),source=self.source.text().strip(),responsible_agency=self.agency.text().strip(),
                    status=self.status.currentText(),confidentiality=self.conf.currentText(),location_text=self.location.text().strip(),
                    lat=number(self.lat.text()),lon=number(self.lon.text()),due_date=self.due.text().strip() or None,
                    mirror_to_neighborhood=self.mirror.isChecked())


class ReferralDialog(QDialog):
    def __init__(self, db, zone_id, parent=None):
        super().__init__(parent); self.setWindowTitle("ارجاع مسئله به کمیته"); self.resize(620,400); form=_dialog_form(self)
        self.issue=QComboBox();
        for x in db.get_social_issues(zone_id):self.issue.addItem(f"{x['title']} — {x['urgency']}",x['id'])
        self.committee=QComboBox();
        for c in db.get_zone_committees(zone_id,ensure=True):self.committee.addItem(c["title"],c["id"])
        self.note=QTextEdit(); self.status=QComboBox();self.status.addItems(["ارجاع‌شده","دریافت‌شده","در حال بررسی","پاسخ‌داده‌شده","مختومه"])
        form.addRow("مسئله:",self.issue);form.addRow("کمیته:",self.committee);form.addRow("شرح ارجاع:",self.note);form.addRow("وضعیت:",self.status)
        buttons=QHBoxLayout();save=set_button_style(QPushButton("ثبت ارجاع"),"link","primary");cancel=QPushButton("انصراف");save.clicked.connect(self.accept);cancel.clicked.connect(self.reject);buttons.addStretch();buttons.addWidget(cancel);buttons.addWidget(save);form.addRow(buttons)
    def data(self):return self.issue.currentData(),self.committee.currentData(),self.note.toPlainText().strip(),self.status.currentText()


class ReferralResponseDialog(QDialog):
    def __init__(self,item,parent=None):
        super().__init__(parent);self.setWindowTitle("ثبت پاسخ کمیته");self.resize(560,360);form=_dialog_form(self)
        self.response=QTextEdit(item.get("response_text") or "");self.status=QComboBox();self.status.addItems(["ارجاع‌شده","دریافت‌شده","در حال بررسی","پاسخ‌داده‌شده","مختومه"]);self.status.setCurrentText(item.get("status") or "پاسخ‌داده‌شده")
        form.addRow("پاسخ/نتیجه:",self.response);form.addRow("وضعیت:",self.status);b=QPushButton("ذخیره");b.clicked.connect(self.accept);form.addRow(b)


class ResolutionDialog(QDialog):
    def __init__(self,db,zone_id,item=None,parent=None):
        super().__init__(parent);self.item=item or {};self.setWindowTitle("مصوبه شورای اجتماعی");self.resize(660,620);form=_dialog_form(self)
        self.meeting=QComboBox();self.meeting.addItem("— بدون جلسه —",None)
        for x in db.get_social_meetings(zone_id):self.meeting.addItem(f"{x.get('meeting_date') or ''} — {x['title']}",x['id'])
        self.meeting.setCurrentIndex(max(0,self.meeting.findData(self.item.get("meeting_id"))))
        self.issue=QComboBox();self.issue.addItem("— بدون مسئله —",None)
        for x in db.get_social_issues(zone_id):self.issue.addItem(x["title"],x["id"])
        self.issue.setCurrentIndex(max(0,self.issue.findData(self.item.get("issue_id"))))
        self.title=QLineEdit(self.item.get("title") or "");self.description=QTextEdit(self.item.get("description") or "")
        self.agency=QLineEdit(self.item.get("responsible_agency") or "");self.person=QLineEdit(self.item.get("responsible_person") or "")
        self.due=QLineEdit(self.item.get("due_date") or "");self.status=QComboBox();self.status.addItems(["در انتظار اقدام","در حال پیگیری","انجام‌شده","لغوشده"]);self.status.setCurrentText(self.item.get("status") or "در انتظار اقدام")
        for label,w in [("جلسه:",self.meeting),("مسئله مرتبط:",self.issue),("عنوان:",self.title),("شرح:",self.description),("دستگاه مسئول:",self.agency),("مسئول پیگیری:",self.person),("مهلت:",self.due),("وضعیت:",self.status)]:form.addRow(label,w)
        b=QPushButton("ذخیره");b.clicked.connect(self.accept);form.addRow(b)
    def data(self):return dict(meeting_id=self.meeting.currentData(),issue_id=self.issue.currentData(),title=self.title.text().strip(),description=self.description.toPlainText().strip(),responsible_agency=self.agency.text().strip(),responsible_person=self.person.text().strip(),due_date=self.due.text().strip() or None,status=self.status.currentText())


class ActionDialog(QDialog):
    def __init__(self,db,zone_id,item=None,parent=None):
        super().__init__(parent);self.item=item or {};self.setWindowTitle("برنامه اقدام اجتماعی");self.resize(700,720);form=_dialog_form(self)
        self.resolution=QComboBox();self.resolution.addItem("— بدون مصوبه —",None)
        for x in db.get_social_resolutions(zone_id):self.resolution.addItem(x["title"],x["id"])
        self.resolution.setCurrentIndex(max(0,self.resolution.findData(self.item.get("resolution_id"))))
        self.issue=QComboBox();self.issue.addItem("— بدون مسئله —",None)
        for x in db.get_social_issues(zone_id):self.issue.addItem(x["title"],x["id"])
        self.issue.setCurrentIndex(max(0,self.issue.findData(self.item.get("issue_id"))))
        self.title=QLineEdit(self.item.get("title") or "");self.description=QTextEdit(self.item.get("action_description") or "")
        self.person=QLineEdit(self.item.get("responsible_person") or "");self.agency=QLineEdit(self.item.get("responsible_agency") or "")
        self.partners=QLineEdit(self.item.get("partner_agencies") or "");self.resources=QTextEdit(self.item.get("required_resources") or "")
        self.budget=QDoubleSpinBox();self.budget.setRange(0,1e15);self.budget.setDecimals(0);self.budget.setValue(float(self.item.get("budget_amount") or 0))
        self.funding=QLineEdit(self.item.get("funding_source") or "");self.start=QLineEdit(self.item.get("start_date") or "");self.end=QLineEdit(self.item.get("end_date") or "")
        self.progress=QSpinBox();self.progress.setRange(0,100);self.progress.setSuffix("٪");self.progress.setValue(int(self.item.get("progress_percent") or 0))
        self.status=QComboBox();self.status.addItems(["برنامه‌ریزی‌شده","در حال اجرا","متوقف","تکمیل‌شده","لغوشده"]);self.status.setCurrentText(self.item.get("status") or "برنامه‌ریزی‌شده")
        self.delay=QTextEdit(self.item.get("delay_reason") or "");self.result=QTextEdit(self.item.get("final_result") or "")
        for label,w in [("مصوبه:",self.resolution),("مسئله:",self.issue),("عنوان اقدام:",self.title),("شرح اقدام:",self.description),("مسئول اجرا:",self.person),("دستگاه مسئول:",self.agency),("دستگاه‌های همکار:",self.partners),("منابع موردنیاز:",self.resources),("بودجه:",self.budget),("منبع اعتبار:",self.funding),("شروع:",self.start),("پایان:",self.end),("پیشرفت:",self.progress),("وضعیت:",self.status),("علت تأخیر:",self.delay),("نتیجه نهایی:",self.result)]:form.addRow(label,w)
        b=QPushButton("ذخیره");b.clicked.connect(self.accept);form.addRow(b)
    def data(self):return dict(resolution_id=self.resolution.currentData(),issue_id=self.issue.currentData(),title=self.title.text().strip(),action_description=self.description.toPlainText().strip(),responsible_person=self.person.text().strip(),responsible_agency=self.agency.text().strip(),partner_agencies=self.partners.text().strip(),required_resources=self.resources.toPlainText().strip(),budget_amount=self.budget.value(),funding_source=self.funding.text().strip(),start_date=self.start.text().strip() or None,end_date=self.end.text().strip() or None,progress_percent=self.progress.value(),status=self.status.currentText(),delay_reason=self.delay.toPlainText().strip(),final_result=self.result.toPlainText().strip())


class SocialCouncilWindow(QWidget):
    back_requested = pyqtSignal()
    def __init__(self,db):
        super().__init__();self.db=db;self.zone_id=None;self.setWindowTitle("شورای اجتماعی محله");self.resize(1500,920);self._build();self.refresh_zones()

    def _build(self):
        root=QVBoxLayout(self);root.addWidget(build_official_header("مدیریت جامع شورای اجتماعی محله",self.db))
        top=QHBoxLayout();back=set_button_style(QPushButton("بازگشت"),"back","ghost");back.clicked.connect(self.back_requested.emit);top.addWidget(back)
        top.addWidget(QLabel("بلوک:"));self.zone=QComboBox();self.zone.setMinimumWidth(300);self.zone.currentIndexChanged.connect(self._zone_changed);top.addWidget(self.zone);sync=QPushButton("همگام‌سازی اعضا و نمایندگان");sync.clicked.connect(self.sync_members);top.addWidget(sync);top.addStretch();root.addLayout(top)
        self.tabs=QTabWidget();root.addWidget(self.tabs,1)
        self._build_dashboard_tab();self._build_members_tab();self._build_meetings_tab();self._build_issues_tab();self._build_referrals_tab();self._build_resolutions_tab();self._build_actions_tab();self._build_chart_reports_tab();self._build_reports_tab()

    def _build_dashboard_tab(self):
        page=QWidget();l=QVBoxLayout(page);profile=QGroupBox("پرونده شورای اجتماعی بلوک");form=QFormLayout(profile)
        self.council_title=QLineEdit();self.formation_date=QLineEdit();self.chair=QComboBox();self.secretary=QComboBox();self.council_status=QComboBox();self.council_status.addItems(["فعال","غیرفعال","تعلیق"]);self.council_notes=QTextEdit();self.council_notes.setMaximumHeight(90)
        for label,w in [("عنوان:",self.council_title),("تاریخ تشکیل:",self.formation_date),("رئیس:",self.chair),("دبیر:",self.secretary),("وضعیت:",self.council_status),("توضیحات:",self.council_notes)]:form.addRow(label,w)
        save=QPushButton("ذخیره پرونده شورا");save.clicked.connect(self.save_profile);form.addRow(save);l.addWidget(profile)
        metrics=QGridLayout();self.metric_labels={}
        specs=[("members_count","اعضای فعال"),("meetings_count","جلسات"),("open_issues","مسائل باز"),("critical_issues","مسائل بحرانی"),("referrals_open","ارجاعات باز"),("pending_resolutions","مصوبات باز"),("actions_active","اقدامات فعال"),("average_progress","میانگین پیشرفت"),("confidential_cases","پرونده محرمانه")]
        for i,(key,title) in enumerate(specs):
            box=QFrame();box.setObjectName("MetricCard");bl=QVBoxLayout(box);v=QLabel("۰");v.setAlignment(Qt.AlignCenter);v.setStyleSheet("font-size:24px;font-weight:700;color:#12396b");t=QLabel(title);t.setAlignment(Qt.AlignCenter);bl.addWidget(v);bl.addWidget(t);self.metric_labels[key]=v;metrics.addWidget(box,i//3,i%3)
        l.addLayout(metrics);self.alerts=QTextEdit();self.alerts.setReadOnly(True);self.alerts.setMinimumHeight(180);l.addWidget(QLabel("هشدارهای مدیریتی"));l.addWidget(self.alerts);l.addStretch();self.tabs.addTab(_scroll(page),get_icon("home","navy"),"داشبورد اجتماعی")

    def _build_members_tab(self):
        page=QWidget();l=QVBoxLayout(page);row=QHBoxLayout();
        for text,slot in [("ثبت عضو جدید",self.add_member),("ویرایش عضو",self.edit_member),("حذف عضو",self.delete_member),("همگام‌سازی خودکار",self.sync_members)]:b=QPushButton(text);b.clicked.connect(slot);row.addWidget(b)
        row.addStretch();l.addLayout(row);self.members=_table(["نام","سمت","نوع نمایندگی","کمیته","دستگاه","تلفن","وضعیت"]);l.addWidget(self.members);self.tabs.addTab(page,get_icon("users","navy"),"اعضای شورا")

    def _build_meetings_tab(self):
        page=QWidget();l=QVBoxLayout(page);row=QHBoxLayout();
        for text,slot in [("ثبت جلسه",self.add_meeting),("ویرایش جلسه",self.edit_meeting),("حذف جلسه",self.delete_meeting)]:b=QPushButton(text);b.clicked.connect(slot);row.addWidget(b)
        row.addStretch();l.addLayout(row);self.meetings=_table(["عنوان","تاریخ","ساعت","محل","وضعیت","دستور جلسه","مدعوین"]);l.addWidget(self.meetings);self.tabs.addTab(page,get_icon("calendar","navy"),"جلسات")

    def _build_issues_tab(self):
        page=QWidget();l=QVBoxLayout(page);row=QHBoxLayout();
        for text,slot in [("ثبت مسئله اجتماعی",self.add_issue),("ویرایش مسئله",self.edit_issue),("حذف مسئله",self.delete_issue),("ارجاع به کمیته",self.add_referral)]:b=QPushButton(text);b.clicked.connect(slot);row.addWidget(b)
        row.addStretch();l.addLayout(row);self.issues=_table(["عنوان","دسته","فوریت","گروه هدف","افراد/خانوار","دستگاه مسئول","وضعیت","محرمانگی","محل"]);l.addWidget(self.issues);self.tabs.addTab(page,get_icon("warning","navy"),"مسائل اجتماعی")

    def _build_referrals_tab(self):
        page=QWidget();l=QVBoxLayout(page);row=QHBoxLayout();add=QPushButton("ثبت ارجاع");add.clicked.connect(self.add_referral);edit=QPushButton("ثبت پاسخ/تغییر وضعیت");edit.clicked.connect(self.edit_referral);delete=QPushButton("حذف ارجاع");delete.clicked.connect(self.delete_referral);row.addWidget(add);row.addWidget(edit);row.addWidget(delete);row.addStretch();l.addLayout(row)
        self.referrals=_table(["مسئله","کمیته","تاریخ ارجاع","فوریت","وضعیت","شرح ارجاع","پاسخ کمیته"]);l.addWidget(self.referrals);self.tabs.addTab(page,get_icon("link","navy"),"ارجاعات به کمیته‌ها")

    def _build_resolutions_tab(self):
        page=QWidget();l=QVBoxLayout(page);row=QHBoxLayout();
        for text,slot in [("ثبت مصوبه",self.add_resolution),("ویرایش مصوبه",self.edit_resolution),("حذف مصوبه",self.delete_resolution),("ساخت برنامه اقدام",self.add_action_from_resolution)]:b=QPushButton(text);b.clicked.connect(slot);row.addWidget(b)
        row.addStretch();l.addLayout(row);self.resolutions=_table(["عنوان","جلسه","مسئله","دستگاه مسئول","مسئول پیگیری","مهلت","وضعیت"]);l.addWidget(self.resolutions);self.tabs.addTab(page,get_icon("resolution","navy"),"مصوبات")

    def _build_actions_tab(self):
        page=QWidget();l=QVBoxLayout(page);row=QHBoxLayout();
        for text,slot in [("ثبت برنامه عملیاتی",self.add_action),("ویرایش برنامه",self.edit_action),("حذف برنامه",self.delete_action)]:b=QPushButton(text);b.clicked.connect(slot);row.addWidget(b)
        row.addStretch();l.addLayout(row);self.actions=_table(["عنوان","مسئله/مصوبه","مسئول","دستگاه","شروع","پایان","پیشرفت","وضعیت","بودجه"]);l.addWidget(self.actions);self.tabs.addTab(page,get_icon("check","navy"),"برنامه‌های عملیاتی")

    def _build_chart_reports_tab(self):
        page=QWidget();layout=QVBoxLayout(page)
        filters=QGroupBox("فیلتر گزارش‌های نموداری");grid=QGridLayout(filters)
        self.chart_type=QComboBox()
        for key,title in CHART_TYPES:self.chart_type.addItem(title,key)
        self.chart_zone=QComboBox();self.chart_committee=QComboBox()
        self.chart_status=QComboBox();self.chart_status.addItem("همه وضعیت‌ها","")
        self.chart_status.addItems(["ثبت اولیه","در حال بررسی","ارجاع‌شده","در حال اقدام","مختومه","در انتظار اقدام","در حال پیگیری","انجام‌شده","برنامه‌ریزی‌شده","در حال اجرا","متوقف","تکمیل‌شده","لغوشده","دریافت‌شده","پاسخ‌داده‌شده"])
        self.chart_date_from=QLineEdit();self.chart_date_from.setPlaceholderText("مثال: ۱۴۰۵/۰۱/۰۱ یا 2026-03-21")
        self.chart_date_to=QLineEdit();self.chart_date_to.setPlaceholderText("مثال: ۱۴۰۵/۱۲/۲۹ یا 2027-03-20")
        grid.addWidget(QLabel("نوع نمودار:"),0,0);grid.addWidget(self.chart_type,0,1)
        grid.addWidget(QLabel("بلوک:"),0,2);grid.addWidget(self.chart_zone,0,3)
        grid.addWidget(QLabel("کمیته:"),1,0);grid.addWidget(self.chart_committee,1,1)
        grid.addWidget(QLabel("وضعیت:"),1,2);grid.addWidget(self.chart_status,1,3)
        grid.addWidget(QLabel("از تاریخ:"),2,0);grid.addWidget(self.chart_date_from,2,1)
        grid.addWidget(QLabel("تا تاریخ:"),2,2);grid.addWidget(self.chart_date_to,2,3)
        buttons=QHBoxLayout();refresh=QPushButton("به‌روزرسانی نمودار");refresh.clicked.connect(self.refresh_chart_report)
        image=QPushButton("ذخیره تصویر نمودار");image.clicked.connect(self.export_chart_image)
        word=QPushButton("خروجی Word همه نمودارها");word.clicked.connect(self.export_charts_word)
        pdf=QPushButton("خروجی PDF همه نمودارها");pdf.clicked.connect(self.export_charts_pdf)
        buttons.addWidget(refresh);buttons.addWidget(image);buttons.addWidget(word);buttons.addWidget(pdf);buttons.addStretch()
        grid.addLayout(buttons,3,0,1,4);layout.addWidget(filters)
        self.chart_widget=BarChartWidget();layout.addWidget(_scroll(self.chart_widget),1)
        self.chart_table=_table(["عنوان","مقدار"]);self.chart_table.setMaximumHeight(230);layout.addWidget(self.chart_table)
        self.chart_type.currentIndexChanged.connect(self.refresh_chart_report)
        self.chart_zone.currentIndexChanged.connect(self._chart_zone_filter_changed)
        self.chart_committee.currentIndexChanged.connect(self.refresh_chart_report)
        self.chart_status.currentIndexChanged.connect(self.refresh_chart_report)
        self.tabs.addTab(page,get_icon("report","navy"),"گزارش‌های نموداری")

    def _build_reports_tab(self):
        page=QWidget();l=QVBoxLayout(page);info=QLabel("گزارش‌ها با رعایت سطح محرمانگی تولید می‌شوند. در خروجی عمومی، جزئیات پرونده‌های محرمانه حذف می‌شود.");info.setWordWrap(True);l.addWidget(info)
        row=QHBoxLayout();
        for text,slot in [("خروجی Excel",self.export_excel),("خروجی Word",self.export_word),("خروجی PDF",self.export_pdf),("گزارش مقایسه بلوک‌ها",self.export_city_excel)]:b=QPushButton(text);b.clicked.connect(slot);row.addWidget(b)
        row.addStretch();l.addLayout(row);self.report_preview=QTextEdit();self.report_preview.setReadOnly(True);l.addWidget(self.report_preview,1);self.tabs.addTab(page,get_icon("report","navy"),"گزارش‌ها")

    def refresh_zones(self):
        zones=self.db.get_zones();current=self.zone.currentData();self.zone.blockSignals(True);self.zone.clear()
        for z in zones:self.zone.addItem(z["name"],z["id"])
        idx=self.zone.findData(current);self.zone.setCurrentIndex(idx if idx>=0 else (0 if self.zone.count() else -1));self.zone.blockSignals(False)
        if hasattr(self,"chart_zone"):
            chart_current=self.chart_zone.currentData();self.chart_zone.blockSignals(True);self.chart_zone.clear();self.chart_zone.addItem("همه بلوک‌ها",None)
            for z in zones:self.chart_zone.addItem(z["name"],z["id"])
            chart_idx=self.chart_zone.findData(chart_current);self.chart_zone.setCurrentIndex(chart_idx if chart_idx>=0 else 0);self.chart_zone.blockSignals(False)
            self._refresh_chart_committee_filter()
        self._zone_changed()

    def _zone_changed(self):
        self.zone_id=self.zone.currentData()
        if not self.zone_id:return
        self.db.ensure_social_council(self.zone_id);self.refresh_all()

    def _can_view_issue(self, item):
        role = (self.db.get_current_user() or {}).get("role")
        level = (item or {}).get("confidentiality") or (item or {}).get("issue_confidentiality") or "داخلی"
        if role == "admin":
            return True
        if role == "manager":
            return level != "فقط مدیر سیستم"
        return level in ("عمومی", "داخلی")

    def refresh_all(self):
        if not self.zone_id:return
        self._refresh_profile();self._refresh_members();self._refresh_meetings();self._refresh_issues();self._refresh_referrals();self._refresh_resolutions();self._refresh_actions();self._refresh_dashboard();self._refresh_report_preview();self.refresh_chart_report()

    def _refresh_profile(self):
        x=self.db.ensure_social_council(self.zone_id) or {};self.council_title.setText(x.get("title") or "");self.formation_date.setText(x.get("formation_date") or "");self.council_status.setCurrentText(x.get("status") or "فعال");self.council_notes.setPlainText(x.get("notes") or "")
        members=self.db.get_social_council_members(self.zone_id,active_only=True);self.chair.clear();self.secretary.clear();self.chair.addItem("— انتخاب —",None);self.secretary.addItem("— انتخاب —",None)
        for m in members:self.chair.addItem(m["full_name"],m["id"]);self.secretary.addItem(m["full_name"],m["id"])
        self.chair.setCurrentIndex(max(0,self.chair.findData(x.get("chair_member_id"))));self.secretary.setCurrentIndex(max(0,self.secretary.findData(x.get("secretary_member_id"))))

    def save_profile(self):
        self.db.update_social_council(self.zone_id,title=self.council_title.text().strip(),formation_date=self.formation_date.text().strip() or None,chair_member_id=self.chair.currentData(),secretary_member_id=self.secretary.currentData(),status=self.council_status.currentText(),notes=self.council_notes.toPlainText().strip());QMessageBox.information(self,"ثبت شد","پرونده شورای اجتماعی ذخیره شد.");self.refresh_all()

    def sync_members(self):
        if not self.zone_id:return
        self.db.sync_social_council_members(self.zone_id);self.refresh_all();QMessageBox.information(self,"همگام‌سازی","معتمدان، مسئولان اماکن و نمایندگان کمیته‌ها بدون حذف اعضای دستی همگام شدند.")

    def _refresh_members(self):
        rows=self.db.get_social_council_members(self.zone_id);self.members.setRowCount(len(rows))
        for r,x in enumerate(rows):
            vals=[x["full_name"],x["role_title"],x["representation_type"],x.get("committee_title"),x.get("agency_name"),x.get("mobile"),x.get("status")]
            for c,v in enumerate(vals):self.members.setItem(r,c,_item(v,x["id"] if c==0 else None))
    def add_member(self):
        d=MemberDialog(self.db,self.zone_id,parent=self)
        if d.exec_():
            try:self.db.add_social_council_member(self.zone_id,**d.data());self.refresh_all()
            except Exception as e:QMessageBox.warning(self,"خطا",str(e))
    def edit_member(self):
        mid=_selected_id(self.members);x=self.db.get_social_council_member(mid) if mid else None
        if not x:return
        d=MemberDialog(self.db,self.zone_id,x,self)
        if d.exec_():self.db.update_social_council_member(mid,**d.data());self.refresh_all()
    def delete_member(self):
        mid=_selected_id(self.members)
        if mid and QMessageBox.question(self,"حذف","عضو انتخاب‌شده حذف شود؟")==QMessageBox.Yes:self.db.delete_social_council_member(mid);self.refresh_all()

    def _refresh_meetings(self):
        rows=self.db.get_social_meetings(self.zone_id);self.meetings.setRowCount(len(rows))
        for r,x in enumerate(rows):
            vals=[x["title"],x.get("meeting_date"),x.get("start_time"),x.get("place_name"),x.get("status"),x.get("agenda"),x.get("invitees")]
            for c,v in enumerate(vals):self.meetings.setItem(r,c,_item(v,x["id"] if c==0 else None))
    def add_meeting(self):
        d=MeetingDialog(self.db,self.zone_id,parent=self)
        if d.exec_():self.db.add_social_meeting(self.zone_id,**d.data());self.refresh_all()
    def edit_meeting(self):
        mid=_selected_id(self.meetings);x=self.db.get_social_meeting(mid) if mid else None
        if x:
            d=MeetingDialog(self.db,self.zone_id,x,self)
            if d.exec_():self.db.update_social_meeting(mid,**d.data());self.refresh_all()
    def delete_meeting(self):
        mid=_selected_id(self.meetings)
        if mid and QMessageBox.question(self,"حذف","جلسه حذف شود؟")==QMessageBox.Yes:self.db.delete_social_meeting(mid);self.refresh_all()

    def _refresh_issues(self):
        rows=[x for x in self.db.get_social_issues(self.zone_id) if self._can_view_issue(x)];self.issues.setRowCount(len(rows))
        for r,x in enumerate(rows):
            vals=[x["title"],x["category"],x["urgency"],x["target_group"],f"{x['affected_people']} / {x['affected_households']}",x.get("responsible_agency"),x["status"],x["confidentiality"],x.get("location_text")]
            for c,v in enumerate(vals):self.issues.setItem(r,c,_item(v,x["id"] if c==0 else None))
    def add_issue(self):
        d=IssueDialog(self.db,parent=self)
        if d.exec_():
            data=d.data();mirror=data.pop("mirror_to_neighborhood",True);self.db.add_social_issue(self.zone_id,mirror_to_neighborhood=mirror,**data);self.refresh_all()
    def edit_issue(self):
        iid=_selected_id(self.issues);x=self.db.get_social_issue(iid) if iid else None
        if x:
            d=IssueDialog(self.db,x,self)
            if d.exec_():data=d.data();data.pop("mirror_to_neighborhood",None);self.db.update_social_issue(iid,**data);self.refresh_all()
    def delete_issue(self):
        iid=_selected_id(self.issues)
        if iid and QMessageBox.question(self,"حذف","پرونده اجتماعی و پیوند عمومی آن حذف شود؟")==QMessageBox.Yes:self.db.delete_social_issue(iid);self.refresh_all()

    def _refresh_referrals(self):
        rows=[x for x in self.db.get_social_referrals(self.zone_id) if self._can_view_issue(x)];self._ref_rows=rows;self.referrals.setRowCount(len(rows))
        for r,x in enumerate(rows):
            vals=[x["issue_title"],x["committee_title"],x.get("referral_date"),x.get("urgency"),x.get("status"),x.get("referral_note"),x.get("response_text")]
            for c,v in enumerate(vals):self.referrals.setItem(r,c,_item(v,(x["issue_id"],x["committee_id"]) if c==0 else None))
    def add_referral(self):
        if not self.db.get_social_issues(self.zone_id):QMessageBox.information(self,"ارجاع","ابتدا یک مسئله اجتماعی ثبت کنید.");return
        d=ReferralDialog(self.db,self.zone_id,self)
        if d.exec_():self.db.refer_social_issue(*d.data());self.refresh_all()
    def edit_referral(self):
        key=_selected_id(self.referrals)
        if not key:return
        x=next((r for r in self._ref_rows if (r["issue_id"],r["committee_id"])==tuple(key)),None)
        if not x:return
        d=ReferralResponseDialog(x,self)
        if d.exec_():self.db.update_social_referral(key[0],key[1],d.response.toPlainText().strip(),d.status.currentText());self.refresh_all()
    def delete_referral(self):
        key=_selected_id(self.referrals)
        if key:self.db.delete_social_referral(key[0],key[1]);self.refresh_all()

    def _refresh_resolutions(self):
        rows=[x for x in self.db.get_social_resolutions(self.zone_id) if self._can_view_issue(x)];self.resolutions.setRowCount(len(rows))
        for r,x in enumerate(rows):
            vals=[x["title"],x.get("meeting_title"),x.get("issue_title"),x.get("responsible_agency"),x.get("responsible_person"),x.get("due_date"),x.get("status")]
            for c,v in enumerate(vals):self.resolutions.setItem(r,c,_item(v,x["id"] if c==0 else None))
    def add_resolution(self):
        d=ResolutionDialog(self.db,self.zone_id,parent=self)
        if d.exec_():self.db.add_social_resolution(self.zone_id,**d.data());self.refresh_all()
    def edit_resolution(self):
        rid=_selected_id(self.resolutions);x=self.db.get_social_resolution(rid) if rid else None
        if x:
            d=ResolutionDialog(self.db,self.zone_id,x,self)
            if d.exec_():self.db.update_social_resolution(rid,**d.data());self.refresh_all()
    def delete_resolution(self):
        rid=_selected_id(self.resolutions)
        if rid and QMessageBox.question(self,"حذف","مصوبه حذف شود؟")==QMessageBox.Yes:self.db.delete_social_resolution(rid);self.refresh_all()
    def add_action_from_resolution(self):
        rid=_selected_id(self.resolutions);x=self.db.get_social_resolution(rid) if rid else None
        d=ActionDialog(self.db,self.zone_id,{"resolution_id":rid,"issue_id":(x or {}).get("issue_id"),"title":f"اقدام: {(x or {}).get('title','')}"},self)
        if d.exec_():self.db.add_social_action_plan(self.zone_id,**d.data());self.refresh_all();self.tabs.setCurrentIndex(6)

    def _refresh_actions(self):
        rows=[x for x in self.db.get_social_action_plans(self.zone_id) if self._can_view_issue(x)];self.actions.setRowCount(len(rows))
        for r,x in enumerate(rows):
            link=x.get("issue_title") or x.get("resolution_title") or "";vals=[x["title"],link,x.get("responsible_person"),x.get("responsible_agency"),x.get("start_date"),x.get("end_date"),f"{x.get('progress_percent') or 0}٪",x.get("status"),f"{float(x.get('budget_amount') or 0):,.0f}"]
            for c,v in enumerate(vals):self.actions.setItem(r,c,_item(v,x["id"] if c==0 else None))
    def add_action(self):
        d=ActionDialog(self.db,self.zone_id,parent=self)
        if d.exec_():self.db.add_social_action_plan(self.zone_id,**d.data());self.refresh_all()
    def edit_action(self):
        aid=_selected_id(self.actions);x=self.db.get_social_action_plan(aid) if aid else None
        if x:
            d=ActionDialog(self.db,self.zone_id,x,self)
            if d.exec_():self.db.update_social_action_plan(aid,**d.data());self.refresh_all()
    def delete_action(self):
        aid=_selected_id(self.actions)
        if aid and QMessageBox.question(self,"حذف","برنامه عملیاتی حذف شود؟")==QMessageBox.Yes:self.db.delete_social_action_plan(aid);self.refresh_all()

    def _refresh_dashboard(self):
        data=self.db.get_social_dashboard(self.zone_id)
        for k,l in self.metric_labels.items():
            value=data.get(k,0);l.setText(to_persian_digits(f"{value}٪" if k=="average_progress" else value))
        alerts=[]
        if data["critical_issues"]:alerts.append(f"• {data['critical_issues']} مسئله فوری یا بحرانی باز است.")
        if data["pending_resolutions"]:alerts.append(f"• {data['pending_resolutions']} مصوبه هنوز تعیین تکلیف نشده است.")
        if data["referrals_open"]:alerts.append(f"• {data['referrals_open']} ارجاع کمیته در انتظار پاسخ است.")
        for a in self.db.get_social_action_plans(self.zone_id):
            if a.get("status") in ("برنامه‌ریزی‌شده","در حال اجرا") and a.get("end_date"):
                alerts.append(f"• سررسید اقدام «{a['title']}»: {a['end_date']}")
        self.alerts.setPlainText("\n".join(alerts) if alerts else "هشدار فعالی وجود ندارد.")

    def _chart_zone_filter_changed(self):
        self._refresh_chart_committee_filter();self.refresh_chart_report()

    def _refresh_chart_committee_filter(self):
        if not hasattr(self,"chart_committee"):return
        current=self.chart_committee.currentData();zone_id=self.chart_zone.currentData() if hasattr(self,"chart_zone") else None
        titles=[]
        zones=[{"id":zone_id}] if zone_id else self.db.get_zones()
        for zone in zones:
            for committee in self.db.get_zone_committees(zone["id"],ensure=True):
                title=committee.get("title") or ""
                if title and title not in titles:titles.append(title)
        self.chart_committee.blockSignals(True);self.chart_committee.clear();self.chart_committee.addItem("همه کمیته‌ها","")
        for title in titles:self.chart_committee.addItem(title,title)
        idx=self.chart_committee.findData(current);self.chart_committee.setCurrentIndex(idx if idx>=0 else 0);self.chart_committee.blockSignals(False)

    def _chart_filter_values(self):
        return {
            "zone_id":self.chart_zone.currentData() if hasattr(self,"chart_zone") else self.zone_id,
            "committee":self.chart_committee.currentData() if hasattr(self,"chart_committee") else "",
            "status":("" if self.chart_status.currentIndex()==0 else self.chart_status.currentText()) if hasattr(self,"chart_status") else "",
            "date_from":self.chart_date_from.text().strip() if hasattr(self,"chart_date_from") else "",
            "date_to":self.chart_date_to.text().strip() if hasattr(self,"chart_date_to") else "",
        }

    def _chart_subtitle(self,filters):
        parts=[]
        zone_id=filters.get("zone_id")
        if zone_id:
            zone=self.db.get_zone(zone_id) or {};parts.append(f"بلوک: {zone.get('name','')}")
        else:parts.append("همه بلوک‌ها")
        if filters.get("committee"):parts.append(f"کمیته: {filters['committee']}")
        if filters.get("status"):parts.append(f"وضعیت: {filters['status']}")
        if filters.get("date_from"):parts.append(f"از: {filters['date_from']}")
        if filters.get("date_to"):parts.append(f"تا: {filters['date_to']}")
        return " | ".join(parts)

    def _chart_zone_rows(self,filters):
        zone_id=filters.get("zone_id")
        return [self.db.get_zone(zone_id)] if zone_id and self.db.get_zone(zone_id) else self.db.get_zones()

    def _committee_issue_ids(self,zone_id,committee_title):
        if not committee_title:return None
        return {row.get("issue_id") for row in self.db.get_social_referrals(zone_id) if row.get("committee_title")==committee_title}

    def _chart_payload(self,chart_key=None):
        chart_key=chart_key or self.chart_type.currentData();filters=self._chart_filter_values();subtitle=self._chart_subtitle(filters)
        zones=self._chart_zone_rows(filters);committee=filters.get("committee") or "";status=filters.get("status") or ""
        date_from,date_to=filters.get("date_from") or "",filters.get("date_to") or ""
        if chart_key=="issues_by_category":
            rows=[]
            for zone in zones:
                issue_ids=self._committee_issue_ids(zone["id"],committee)
                items=[x for x in self.db.get_social_issues(zone["id"]) if self._can_view_issue(x)]
                if issue_ids is not None:items=[x for x in items if x.get("id") in issue_ids]
                rows.extend(filter_rows(items,status,date_from,date_to,("created_at","due_date")))
            return issues_by_category_payload(rows,subtitle)
        if chart_key=="blocks_comparison":
            blocks=[]
            for zone in zones:
                issue_ids=self._committee_issue_ids(zone["id"],committee)
                issues=[x for x in self.db.get_social_issues(zone["id"]) if self._can_view_issue(x)]
                resolutions=[x for x in self.db.get_social_resolutions(zone["id"]) if self._can_view_issue(x)]
                actions=[x for x in self.db.get_social_action_plans(zone["id"]) if self._can_view_issue(x)]
                if issue_ids is not None:
                    issues=[x for x in issues if x.get("id") in issue_ids]
                    resolutions=[x for x in resolutions if x.get("issue_id") in issue_ids]
                    actions=[x for x in actions if x.get("issue_id") in issue_ids]
                issues=filter_rows(issues,status,date_from,date_to,("created_at","due_date"))
                resolutions=filter_rows(resolutions,status,date_from,date_to,("created_at","due_date"))
                actions=filter_rows(actions,status,date_from,date_to,("created_at","start_date","end_date"))
                blocks.append({"zone_name":zone.get("name"),"issues":issues,"resolutions":resolutions,"actions":actions})
            return blocks_comparison_payload(blocks,subtitle,exact_status=bool(status))
        if chart_key=="committees_performance":
            refs=[];titles=[]
            for zone in zones:
                for c in self.db.get_zone_committees(zone["id"],ensure=True):
                    title=c.get("title") or ""
                    if title and title not in titles:titles.append(title)
                items=[x for x in self.db.get_social_referrals(zone["id"]) if self._can_view_issue(x)]
                if committee:items=[x for x in items if x.get("committee_title")==committee]
                refs.extend(filter_rows(items,status,date_from,date_to,("referral_date","updated_at")))
            if committee:titles=[committee]
            return committees_performance_payload(refs,titles,subtitle)
        if chart_key=="resolutions_status":
            rows=[]
            for zone in zones:
                issue_ids=self._committee_issue_ids(zone["id"],committee)
                items=[x for x in self.db.get_social_resolutions(zone["id"]) if self._can_view_issue(x)]
                if issue_ids is not None:items=[x for x in items if x.get("issue_id") in issue_ids]
                rows.extend(filter_rows(items,status,date_from,date_to,("created_at","due_date")))
            return resolutions_status_payload(rows,subtitle)
        rows=[]
        for zone in zones:
            issue_ids=self._committee_issue_ids(zone["id"],committee)
            items=[x for x in self.db.get_social_action_plans(zone["id"]) if self._can_view_issue(x)]
            if issue_ids is not None:items=[x for x in items if x.get("issue_id") in issue_ids]
            rows.extend(filter_rows(items,status,date_from,date_to,("created_at","start_date","end_date")))
        return actions_status_payload(rows,subtitle)

    def _apply_chart_payload(self,payload,update_table=True):
        self.chart_widget.set_chart(payload.get("title"),payload.get("categories"),payload.get("series"),payload.get("subtitle"))
        if not update_table:return
        headers=payload.get("headers") or ["عنوان","مقدار"];rows=payload.get("rows") or []
        self.chart_table.clear();self.chart_table.setColumnCount(len(headers));self.chart_table.setHorizontalHeaderLabels(headers);self.chart_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch);self.chart_table.setRowCount(len(rows))
        for r,row in enumerate(rows):
            for c,value in enumerate(row):self.chart_table.setItem(r,c,_item(value))

    def refresh_chart_report(self):
        if not hasattr(self,"chart_widget") or self.chart_zone.count()==0:return
        try:self._apply_chart_payload(self._chart_payload())
        except Exception as exc:
            self.chart_widget.set_chart("گزارش نموداری",[],[],f"خطا در تولید گزارش: {exc}")

    def export_chart_image(self):
        try:
            path,_=QFileDialog.getSaveFileName(self,"ذخیره تصویر نمودار",f"social_chart_{self.chart_type.currentData()}.png","PNG (*.png)")
            if not path:return
            if not path.lower().endswith(".png"):path+=".png"
            self._apply_chart_payload(self._chart_payload())
            if not self.chart_widget.save_png(path):raise RuntimeError("ذخیره تصویر ناموفق بود.")
            QMessageBox.information(self,"خروجی","تصویر نمودار ذخیره شد.")
        except Exception as exc:QMessageBox.warning(self,"خطا",str(exc))

    def _render_all_chart_images(self,directory):
        current=self.chart_type.currentData();results=[]
        for index,(key,title) in enumerate(CHART_TYPES,1):
            payload=self._chart_payload(key);self._apply_chart_payload(payload,False);QApplication.processEvents()
            path=os.path.join(directory,f"chart_{index}_{key}.png")
            if not self.chart_widget.save_png(path):raise RuntimeError(f"ساخت تصویر {title} ناموفق بود.")
            results.append((payload,path))
        idx=self.chart_type.findData(current)
        if idx>=0:self.chart_type.setCurrentIndex(idx)
        self.refresh_chart_report();return results

    def export_charts_word(self):
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.section import WD_ORIENT
            path,_=QFileDialog.getSaveFileName(self,"خروجی Word نمودارها",f"social_bar_charts_{self.zone_id or 'all'}.docx","Word (*.docx)")
            if not path:return
            with tempfile.TemporaryDirectory() as tmp:
                charts=self._render_all_chart_images(tmp);doc=Document();section=doc.sections[0];section.orientation=WD_ORIENT.LANDSCAPE;section.page_width,section.page_height=section.page_height,section.page_width;doc.add_heading("گزارش‌های نموداری شورای اجتماعی",0);doc.add_paragraph(f"تاریخ تولید: {datetime.now():%Y-%m-%d %H:%M}")
                for payload,image_path in charts:
                    doc.add_heading(payload.get("title") or "نمودار",1);doc.add_paragraph(payload.get("subtitle") or "")
                    doc.add_picture(image_path,width=Inches(9.0))
                    headers=payload.get("headers") or [];rows=payload.get("rows") or []
                    if headers:
                        table=doc.add_table(rows=1,cols=len(headers));table.style="Table Grid"
                        for i,value in enumerate(headers):table.rows[0].cells[i].text=_text(value)
                        for row in rows:
                            cells=table.add_row().cells
                            for i,value in enumerate(row):cells[i].text=_text(value)
                doc.save(path)
            QMessageBox.information(self,"خروجی","فایل Word شامل پنج نمودار ساخته شد.")
        except Exception as exc:QMessageBox.warning(self,"خطا",str(exc))

    def export_charts_pdf(self):
        try:
            from reportlab.lib.pagesizes import A4,landscape
            from reportlab.lib.utils import ImageReader
            from reportlab.pdfgen import canvas
            path,_=QFileDialog.getSaveFileName(self,"خروجی PDF نمودارها",f"social_bar_charts_{self.zone_id or 'all'}.pdf","PDF (*.pdf)")
            if not path:return
            with tempfile.TemporaryDirectory() as tmp:
                charts=self._render_all_chart_images(tmp);page=landscape(A4);cv=canvas.Canvas(path,pagesize=page);page_w,page_h=page
                for payload,image_path in charts:
                    image=QPixmap(image_path);iw=max(1,image.width());ih=max(1,image.height());scale=min((page_w-36)/iw,(page_h-36)/ih)
                    draw_w,draw_h=iw*scale,ih*scale;cv.drawImage(ImageReader(image_path),(page_w-draw_w)/2,(page_h-draw_h)/2,draw_w,draw_h,preserveAspectRatio=True,mask="auto");cv.showPage()
                cv.save()
            QMessageBox.information(self,"خروجی","فایل PDF شامل پنج نمودار ساخته شد.")
        except Exception as exc:QMessageBox.warning(self,"خطا",str(exc))

    def _report_data(self,public=False):
        zone=self.db.get_zone(self.zone_id) or {};council=self.db.get_social_council(self.zone_id) or {};stats=self.db.get_social_dashboard(self.zone_id)
        issues=self.db.get_social_issues(self.zone_id,include_confidential=not public)
        if public:issues=[x for x in issues if x.get("confidentiality") in ("عمومی","داخلی")]
        else:issues=[x for x in issues if self._can_view_issue(x)]
        resolutions=[x for x in self.db.get_social_resolutions(self.zone_id) if self._can_view_issue(x)]
        actions=[x for x in self.db.get_social_action_plans(self.zone_id) if self._can_view_issue(x)]
        return zone,council,stats,self.db.get_social_council_members(self.zone_id),self.db.get_social_meetings(self.zone_id),issues,resolutions,actions

    def _refresh_report_preview(self):
        zone,c,stats,members,meetings,issues,resolutions,actions=self._report_data(public=True)
        lines=[f"شناسنامه اجتماعی بلوک: {zone.get('name','')}",f"عنوان شورا: {c.get('title','')}",f"اعضای فعال: {stats['members_count']}",f"جلسات: {stats['meetings_count']}",f"مسائل باز: {stats['open_issues']} | بحرانی: {stats['critical_issues']}",f"مصوبات باز: {stats['pending_resolutions']} | اقدامات فعال: {stats['actions_active']}","","مسائل قابل انتشار:"]
        lines += [f"- {x['title']} | {x['category']} | {x['urgency']} | {x['status']}" for x in issues]
        self.report_preview.setPlainText("\n".join(lines))

    def export_excel(self):
        try:
            from openpyxl import Workbook
            path,_=QFileDialog.getSaveFileName(self,"خروجی Excel",f"social_council_{self.zone_id}.xlsx","Excel (*.xlsx)")
            if not path:return
            zone,c,stats,members,meetings,issues,resolutions,actions=self._report_data(False);wb=Workbook();ws=wb.active;ws.title="خلاصه";ws.append(["شاخص","مقدار"])
            for k,v in stats.items():ws.append([k,v])
            for name,headers,rows in [("اعضا",["نام","سمت","نمایندگی","دستگاه","تلفن","وضعیت"],members),("جلسات",["عنوان","تاریخ","محل","وضعیت"],meetings),("مسائل",["عنوان","دسته","فوریت","گروه هدف","وضعیت","محرمانگی"],issues),("مصوبات",["عنوان","دستگاه","مسئول","مهلت","وضعیت"],resolutions),("برنامه اقدام",["عنوان","مسئول","دستگاه","شروع","پایان","پیشرفت","وضعیت"],actions)]:
                sh=wb.create_sheet(name);sh.append(headers)
                for x in rows:
                    if name=="اعضا":sh.append([x.get("full_name"),x.get("role_title"),x.get("representation_type"),x.get("agency_name"),x.get("mobile"),x.get("status")])
                    elif name=="جلسات":sh.append([x.get("title"),x.get("meeting_date"),x.get("place_name"),x.get("status")])
                    elif name=="مسائل":sh.append([x.get("title"),x.get("category"),x.get("urgency"),x.get("target_group"),x.get("status"),x.get("confidentiality")])
                    elif name=="مصوبات":sh.append([x.get("title"),x.get("responsible_agency"),x.get("responsible_person"),x.get("due_date"),x.get("status")])
                    else:sh.append([x.get("title"),x.get("responsible_person"),x.get("responsible_agency"),x.get("start_date"),x.get("end_date"),x.get("progress_percent"),x.get("status")])
            wb.save(path);QMessageBox.information(self,"خروجی","فایل Excel ساخته شد.")
        except Exception as e:QMessageBox.warning(self,"خطا",str(e))

    def export_city_excel(self):
        try:
            from openpyxl import Workbook
            path,_=QFileDialog.getSaveFileName(self,"مقایسه بلوک‌ها","social_councils_city.xlsx","Excel (*.xlsx)")
            if not path:return
            wb=Workbook();ws=wb.active;ws.title="مقایسه";headers=["بلوک","اعضا","جلسات","مسائل باز","بحرانی","ارجاعات باز","مصوبات باز","اقدامات فعال","میانگین پیشرفت"];ws.append(headers)
            for x in self.db.get_social_city_summary():ws.append([x["zone_name"],x["members_count"],x["meetings_count"],x["open_issues"],x["critical_issues"],x["referrals_open"],x["pending_resolutions"],x["actions_active"],x["average_progress"]])
            wb.save(path);QMessageBox.information(self,"خروجی","گزارش مقایسه ساخته شد.")
        except Exception as e:QMessageBox.warning(self,"خطا",str(e))

    def export_word(self):
        try:
            from docx import Document
            path,_=QFileDialog.getSaveFileName(self,"خروجی Word",f"social_council_{self.zone_id}.docx","Word (*.docx)")
            if not path:return
            zone,c,stats,members,meetings,issues,resolutions,actions=self._report_data(False);doc=Document();doc.add_heading(f"گزارش شورای اجتماعی {zone.get('name','')}",0);doc.add_paragraph(f"تاریخ تولید: {datetime.now():%Y-%m-%d %H:%M}")
            doc.add_heading("خلاصه مدیریتی",1)
            for k,v in stats.items():doc.add_paragraph(f"{k}: {v}")
            for title,rows,fields in [("اعضا",members,["full_name","role_title","agency_name","status"]),("جلسات",meetings,["title","meeting_date","place_name","status"]),("مسائل اجتماعی",issues,["title","category","urgency","target_group","status","confidentiality"]),("مصوبات",resolutions,["title","responsible_agency","responsible_person","due_date","status"]),("برنامه‌های عملیاتی",actions,["title","responsible_agency","progress_percent","status"])]:
                doc.add_heading(title,1)
                for x in rows:doc.add_paragraph(" | ".join(_text(x.get(f)) for f in fields),style="List Bullet")
            doc.save(path);QMessageBox.information(self,"خروجی","فایل Word ساخته شد.")
        except Exception as e:QMessageBox.warning(self,"خطا",str(e))

    def export_pdf(self):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            path,_=QFileDialog.getSaveFileName(self,"خروجی PDF",f"social_council_{self.zone_id}.pdf","PDF (*.pdf)")
            if not path:return
            zone,c,stats,members,meetings,issues,resolutions,actions=self._report_data(True);cv=canvas.Canvas(path,pagesize=A4);w,h=A4;y=h-50
            lines=[f"Social Council Report - {zone.get('name','')}",f"Generated: {datetime.now():%Y-%m-%d %H:%M}",f"Members: {stats['members_count']}  Meetings: {stats['meetings_count']}",f"Open issues: {stats['open_issues']}  Critical: {stats['critical_issues']}","Public/Internal issues:"]+[f"- {x['title']} | {x['category']} | {x['urgency']} | {x['status']}" for x in issues]
            for line in lines:
                if y<50:cv.showPage();y=h-50
                cv.drawString(40,y,line[:110]);y-=18
            cv.save();QMessageBox.information(self,"خروجی","فایل PDF ساخته شد. اطلاعات محرمانه در PDF عمومی حذف شده است.")
        except Exception as e:QMessageBox.warning(self,"خطا",str(e))
