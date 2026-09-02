# -*- coding: utf-8 -*-
"""تولید اسناد رسمی Word و PDF با QR اعتبارسنجی و امضای اسکن‌شده."""

import hashlib
import io
import os
from datetime import datetime
from string import Formatter

from pdf_text_utils import shape_fa
from jalali_utils import convert_dates_in_text, format_jalali, today_jalali

DOCX_AVAILABLE = True
DOCX_IMPORT_ERROR = None
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ModuleNotFoundError as exc:
    DOCX_AVAILABLE = False
    DOCX_IMPORT_ERROR = exc
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    Cm = Pt = RGBColor = None
    OxmlElement = qn = None

QR_AVAILABLE = True
QR_IMPORT_ERROR = None
try:
    import qrcode
except ModuleNotFoundError as exc:
    QR_AVAILABLE = False
    QR_IMPORT_ERROR = exc
    qrcode = None


def ensure_docx_available():
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "کتابخانه ساخت Word نصب نیست. اجرا کنید:\npython3 -m pip install python-docx"
        ) from DOCX_IMPORT_ERROR


def ensure_qr_available():
    if not QR_AVAILABLE:
        raise RuntimeError(
            "کتابخانه QR نصب نیست. اجرا کنید:\npython3 -m pip install qrcode[pil]"
        ) from QR_IMPORT_ERROR


class SafeFormatDict(dict):
    def __missing__(self, key):
        return "—"


def render_template_text(template, context):
    context = SafeFormatDict({k: "—" if v is None or v == "" else v for k, v in (context or {}).items()})
    try:
        return convert_dates_in_text((template or "").format_map(context))
    except (ValueError, KeyError):
        return template or ""


def template_fields(*templates):
    fields = []
    for template in templates:
        try:
            for _, field_name, _, _ in Formatter().parse(template or ""):
                if field_name and field_name not in fields:
                    fields.append(field_name)
        except ValueError:
            continue
    return fields


def _verification_data(subject, body, metadata):
    number = str(metadata.get("number") or "—")
    date = format_jalali(metadata.get("date") or datetime.now().strftime("%Y-%m-%d"))
    recipient = str(metadata.get("recipient") or "—")
    raw = f"{number}|{date}|{recipient}|{subject or ''}|{body or ''}"
    token = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:24].upper()
    base_url = str(metadata.get("verification_base_url") or "").strip().rstrip("/")
    payload = f"{base_url}/{token}" if base_url else f"JAVANROOD-DOC|{token}|{number}|{date}"
    return token, payload


def _qr_png(payload, box_size=7, border=2):
    ensure_qr_available()
    qr = qrcode.QRCode(version=None, error_correction=qrcode.constants.ERROR_CORRECT_M,
                       box_size=box_size, border=border)
    qr.add_data(payload)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    stream = io.BytesIO()
    img.save(stream, format="PNG")
    stream.seek(0)
    return stream

def _official_emblem_path():
    path = os.path.join(os.path.dirname(__file__), "assets", "official_emblem.png")
    return path if os.path.exists(path) else ""


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_borders(table, color="808080", size=8):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_rtl(paragraph, alignment=None):
    ensure_docx_available()
    paragraph.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _style_run(run, size=12, bold=False, color=None, font_name="Tahoma"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_rtl_paragraph(container, text="", size=12, bold=False, alignment=None,
                       color=None, space_after=4, line_spacing=1.25, font_name="Tahoma"):
    ensure_docx_available()
    p = container.add_paragraph()
    _set_paragraph_rtl(p, alignment if alignment is not None else WD_ALIGN_PARAGRAPH.RIGHT)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(convert_dates_in_text(str(text or "")))
    _style_run(run, size=size, bold=bold, color=color, font_name=font_name)
    return p


def _enrich_metadata_from_db(db, metadata):
    enriched = dict(metadata or {})
    try:
        signature = db.get_official_signature()
    except Exception:
        signature = {}
    enriched.setdefault("signature_image_path", signature.get("image_path") or "")
    enriched.setdefault("signature", signature.get("signer_name") or enriched.get("creator") or "")
    enriched.setdefault("signature_title", signature.get("signer_title") or "")
    enriched.setdefault("verification_base_url", signature.get("verification_base_url") or "")
    return enriched


def metadata_from_context(context, db=None, overrides=None):
    context = context or {}
    metadata = {
        "number": context.get("letter_number") or "—",
        "date": format_jalali(context.get("letter_date") or context.get("date") or datetime.now().strftime("%Y-%m-%d")),
        "time": context.get("time") or "—",
        "attachment": context.get("attachment_text") or "ندارد",
        "recipient": context.get("recipient") or "مقام / سازمان مربوطه",
        "sender": context.get("sender") or "",
        "creator": context.get("user_full_name") or "",
    }
    metadata.update({k: v for k, v in (overrides or {}).items() if v not in (None, "")})
    return _enrich_metadata_from_db(db, metadata) if db is not None else metadata


def generate_official_docx(output_path, title, subject, body, metadata=None,
                           organization="فرمانداری شهرستان جوانرود",
                           system_title="سامانه مدیریت محله‌محور"):
    ensure_docx_available()
    metadata = dict(metadata or {})
    token, payload = _verification_data(subject, body, metadata)
    metadata["verification_token"] = token
    metadata["verification_payload"] = payload
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    document.core_properties.title = str(title or subject or "سند اداری")
    document.core_properties.subject = str(subject or "")
    document.core_properties.author = organization

    header = document.add_table(rows=2, cols=3)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    widths = [Cm(4.3), Cm(10.0), Cm(4.3)]
    for i, width in enumerate(widths):
        header.columns[i].width = width
        for r in range(2):
            header.cell(r, i).width = width
            header.cell(r, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(header.cell(r, i), 85, 100, 85, 100)
    _set_table_borders(header, color="666666", size=10)

    left, center, right = header.cell(0, 0), header.cell(0, 1), header.cell(0, 2)
    meta_lines = [
        f"شماره: {metadata.get('number') or '................'}",
        f"تاریخ: {format_jalali(metadata.get('date') or datetime.now().strftime('%Y-%m-%d'))}",
        f"ساعت: {metadata.get('time') or '................'}",
        f"پیوست: {metadata.get('attachment') or 'ندارد'}",
    ]
    for i, line in enumerate(meta_lines):
        p = left.paragraphs[0] if i == 0 else left.add_paragraph()
        _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
        _style_run(p.add_run(line), size=10, bold=(i == 0), font_name="B Nazanin")

    p = center.paragraphs[0]
    _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    _style_run(p.add_run("بسمه تعالی"), size=18, bold=True, font_name="B Titr")
    _add_rtl_paragraph(center, str(title or system_title), size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, color="13294B", font_name="B Nazanin", space_after=0)

    p = right.paragraphs[0]
    _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    emblem_path = _official_emblem_path()
    if emblem_path:
        try:
            p.add_run().add_picture(emblem_path, width=Cm(1.15))
        except Exception:
            pass
    for i, line in enumerate(["جمهوری اسلامی ایران", "وزارت کشور", "استانداری کرمانشاه", organization]):
        p = right.add_paragraph(); _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        _style_run(p.add_run(line), size=9 if i < 3 else 10, bold=True, font_name="B Nazanin")

    merged = header.cell(1, 0).merge(header.cell(1, 1))
    _set_cell_shading(merged, "F6F6F6")
    p = merged.paragraphs[0]; _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    _style_run(p.add_run("«سال اقتصاد مقاومتی در سایه وحدت ملی و امنیت ملی»"), size=10,
               bold=True, font_name="B Nazanin")
    p = header.cell(1, 2).paragraphs[0]; _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    _style_run(p.add_run(system_title), size=9, bold=True, color="13294B", font_name="B Nazanin")

    document.add_paragraph()
    main_table = document.add_table(rows=1, cols=1)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    main_table.columns[0].width = Cm(18.4)
    _set_table_borders(main_table, color="666666", size=12)
    cell = main_table.cell(0, 0)
    _set_cell_margins(cell, 170, 220, 170, 220)

    _add_rtl_paragraph(cell, metadata.get("recipient") or "مقام / سازمان مربوطه", size=14, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_name="B Nazanin", space_after=8)
    _add_rtl_paragraph(cell, metadata.get("greeting") or "سلام علیکم", size=13,
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_name="B Nazanin", space_after=5)
    if subject:
        _add_rtl_paragraph(cell, f"موضوع: {subject}", size=13, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_name="B Nazanin", space_after=7)
    for line in str(body or "").splitlines():
        _add_rtl_paragraph(cell, line, size=13, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           font_name="B Nazanin", line_spacing=1.5, space_after=4)

    document.add_paragraph()
    sign_table = document.add_table(rows=1, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    sign_table.columns[0].width = Cm(9.2)
    sign_table.columns[1].width = Cm(9.2)
    copies = metadata.get("copies") or metadata.get("cc") or []
    if isinstance(copies, str):
        copies = [x.strip() for x in copies.split("\n") if x.strip()]
    if copies:
        _add_rtl_paragraph(sign_table.cell(0, 0), "رونوشت:", size=12, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_name="B Nazanin")
        for i, item in enumerate(copies, 1):
            _add_rtl_paragraph(sign_table.cell(0, 0), f"{i}. {item}", size=11,
                               alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_name="B Nazanin", space_after=1)

    sig_cell = sign_table.cell(0, 1)
    _add_rtl_paragraph(sig_cell, "از طرف", size=12, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="B Nazanin", space_after=2)
    signature_path = metadata.get("signature_image_path") or ""
    if signature_path and os.path.exists(signature_path):
        p = sig_cell.add_paragraph(); _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        try:
            p.add_run().add_picture(signature_path, width=Cm(4.0))
        except Exception:
            pass
    if metadata.get("signature"):
        _add_rtl_paragraph(sig_cell, metadata.get("signature"), size=13, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="B Nazanin", space_after=1)
    if metadata.get("signature_title"):
        _add_rtl_paragraph(sig_cell, metadata.get("signature_title"), size=12, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="B Nazanin", space_after=1)
    _add_rtl_paragraph(sig_cell, "امضاء و مهر", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       font_name="B Nazanin", color="555555", space_after=0)

    document.add_paragraph()
    footer = document.add_table(rows=1, cols=2)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.autofit = False
    footer.columns[0].width = Cm(3.2)
    footer.columns[1].width = Cm(15.2)
    _set_table_borders(footer, color="777777", size=10)
    qr_cell = footer.cell(0, 0)
    p = qr_cell.paragraphs[0]; _set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    try:
        qr_stream = _qr_png(payload, box_size=6)
        p.add_run().add_picture(qr_stream, width=Cm(2.5))
        _add_rtl_paragraph(qr_cell, token, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           font_name="Tahoma", space_after=0)
    except Exception:
        _style_run(p.add_run(token), size=8, bold=True, font_name="Tahoma")

    foot_cell = footer.cell(0, 1)
    _add_rtl_paragraph(foot_cell, metadata.get("footer_address") or
                       "آدرس: کرمانشاه، جوانرود، خیابان ..........، پلاک ..........",
                       size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="B Nazanin", space_after=1)
    _add_rtl_paragraph(foot_cell, metadata.get("footer_website") or "www.ostan.ksh.ir",
                       size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       font_name="Tahoma", space_after=0)

    document.save(output_path)
    return {"path": output_path, "verification_token": token, "verification_payload": payload}


def _pdf_font_setup():
    import report_generator as rg
    rg._ensure_fonts_registered()
    return rg.FONT_NAME, rg.FONT_NAME_BOLD


def _pdf_wrap(text, font_name, font_size, max_width):
    from reportlab.pdfbase.pdfmetrics import stringWidth
    words = str(text or "").split()
    if not words:
        return [""]
    lines, current = [], []
    for word in words:
        trial = " ".join(current + [word])
        if current and stringWidth(shape_fa(trial), font_name, font_size) > max_width:
            lines.append(" ".join(current)); current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


def generate_official_pdf(output_path, title, subject, body, metadata=None,
                          organization="فرمانداری شهرستان جوانرود",
                          system_title="سامانه مدیریت محله‌محور"):
    """تولید PDF A4 با چیدمان هم‌ساختار نسخه Word."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader

    metadata = dict(metadata or {})
    token, payload = _verification_data(subject, body, metadata)
    metadata["verification_token"] = token
    metadata["verification_payload"] = payload
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    regular, bold = _pdf_font_setup()
    page_w, page_h = A4
    c = canvas.Canvas(output_path, pagesize=A4)
    c.setTitle(str(title or subject or "سند اداری"))

    def draw_r(x, y, text, size=10, is_bold=False, color=colors.black):
        c.setFont(bold if is_bold else regular, size)
        c.setFillColor(color)
        c.drawRightString(x, y, shape_fa(text))

    def draw_c(x, y, text, size=10, is_bold=False, color=colors.black):
        c.setFont(bold if is_bold else regular, size)
        c.setFillColor(color)
        c.drawCentredString(x, y, shape_fa(text))

    def header_box():
        margin = 12 * mm
        top = page_h - 12 * mm
        h = 47 * mm
        y = top - h
        c.setStrokeColor(colors.HexColor("#666666")); c.setLineWidth(1.1)
        c.roundRect(margin, y, page_w - 2 * margin, h, 5 * mm, stroke=1, fill=0)
        x1 = margin + 44 * mm; x2 = margin + 145 * mm
        c.line(x1, y, x1, top); c.line(x2, y, x2, top)
        draw_r(x1 - 4 * mm, top - 9 * mm, f"شماره: {metadata.get('number') or '................'}", 9, True)
        draw_r(x1 - 4 * mm, top - 18 * mm, f"تاریخ: {format_jalali(metadata.get('date') or datetime.now().strftime('%Y-%m-%d'))}", 9)
        draw_r(x1 - 4 * mm, top - 27 * mm, f"ساعت: {metadata.get('time') or '................'}", 9)
        draw_r(x1 - 4 * mm, top - 36 * mm, f"پیوست: {metadata.get('attachment') or 'ندارد'}", 9)
        mid = (x1 + x2) / 2
        draw_c(mid, top - 10 * mm, "بسمه تعالی", 15, True)
        draw_c(mid, top - 23 * mm, str(title or system_title), 12, True, colors.HexColor("#13294B"))
        draw_c(mid, top - 36 * mm, "«سال اقتصاد مقاومتی در سایه وحدت ملی و امنیت ملی»", 8.5, True)
        org_mid = (x2 + page_w - margin) / 2
        emblem_path = _official_emblem_path()
        if emblem_path:
            try:
                c.drawImage(ImageReader(emblem_path), org_mid - 7 * mm, top - 11 * mm,
                            width=14 * mm, height=10 * mm, preserveAspectRatio=True, mask='auto')
            except Exception:
                pass
        draw_c(org_mid, top - 16 * mm, "جمهوری اسلامی ایران", 8, True)
        draw_c(org_mid, top - 23 * mm, "وزارت کشور", 8, True)
        draw_c(org_mid, top - 30 * mm, "استانداری کرمانشاه", 8, True)
        draw_c(org_mid, top - 37 * mm, "فرمانداری شهرستان جوانرود", 8, True)
        draw_c(org_mid, top - 43 * mm, system_title, 6.8, True, colors.HexColor("#13294B"))
        return y

    header_bottom = header_box()
    box_x, box_w = 13 * mm, page_w - 26 * mm
    box_top = header_bottom - 8 * mm
    box_bottom = 45 * mm
    c.roundRect(box_x, box_bottom, box_w, box_top - box_bottom, 5 * mm, stroke=1, fill=0)

    y = box_top - 10 * mm
    right_x = box_x + box_w - 8 * mm
    draw_r(right_x, y, metadata.get("recipient") or "مقام / سازمان مربوطه", 12, True); y -= 9 * mm
    draw_r(right_x, y, metadata.get("greeting") or "سلام علیکم", 11); y -= 9 * mm
    if subject:
        draw_r(right_x, y, f"موضوع: {subject}", 11, True); y -= 9 * mm

    max_width = box_w - 18 * mm
    for paragraph in str(body or "").splitlines() or [""]:
        lines = _pdf_wrap(paragraph, regular, 10.5, max_width)
        for line in lines:
            if y < 105 * mm:
                break
            draw_r(right_x, y, line, 10.5)
            y -= 7 * mm
        y -= 2 * mm

    # امضا در پایین سمت چپ بدنه
    sig_x = box_x + 52 * mm
    sig_y = box_bottom + 43 * mm
    draw_c(sig_x, sig_y + 21 * mm, "از طرف", 10, True)
    signature_path = metadata.get("signature_image_path") or ""
    if signature_path and os.path.exists(signature_path):
        try:
            c.drawImage(ImageReader(signature_path), sig_x - 22 * mm, sig_y + 4 * mm,
                        width=44 * mm, height=18 * mm, preserveAspectRatio=True, mask='auto')
        except Exception:
            pass
    if metadata.get("signature"):
        draw_c(sig_x, sig_y, metadata.get("signature"), 11, True)
    if metadata.get("signature_title"):
        draw_c(sig_x, sig_y - 7 * mm, metadata.get("signature_title"), 10, True)
    draw_c(sig_x, sig_y - 14 * mm, "امضاء و مهر", 8)

    copies = metadata.get("copies") or metadata.get("cc") or []
    if isinstance(copies, str):
        copies = [x.strip() for x in copies.split("\n") if x.strip()]
    if copies:
        copy_y = box_bottom + 43 * mm
        draw_r(right_x, copy_y, "رونوشت:", 10, True)
        for i, item in enumerate(copies, 1):
            copy_y -= 7 * mm; draw_r(right_x, copy_y, f"{i}. {item}", 9)

    # فوتر و QR واقعی
    footer_y = 10 * mm
    c.rect(55 * mm, footer_y, 120 * mm, 23 * mm, stroke=1, fill=0)
    draw_c(115 * mm, footer_y + 14 * mm, metadata.get("footer_address") or
           "آدرس: کرمانشاه، جوانرود، خیابان ..........، پلاک ..........", 8.5)
    draw_c(115 * mm, footer_y + 6 * mm, metadata.get("footer_website") or "www.ostan.ksh.ir", 9, True)
    try:
        qr_stream = _qr_png(payload, box_size=7)
        c.drawImage(ImageReader(qr_stream), 14 * mm, footer_y, width=23 * mm, height=23 * mm, mask='auto')
        draw_c(25.5 * mm, footer_y - 3 * mm, token, 5.5)
    except Exception:
        draw_c(25.5 * mm, footer_y + 10 * mm, token, 6, True)

    c.save()
    return {"path": output_path, "verification_token": token, "verification_payload": payload}


def generate_document_from_template(db, template_id, output_path, context=None, zone_id=None,
                                    related_entity_type=None, related_entity_id=None, metadata=None):
    template = db.get_document_template(template_id)
    if not template:
        raise ValueError("قالب انتخاب‌شده پیدا نشد.")
    context = context or {}
    subject = render_template_text(template.get("subject_template") or template["name"], context)
    body = render_template_text(template.get("body_template") or "", context)
    metadata = metadata_from_context(context, db=db, overrides=metadata)
    extension = os.path.splitext(output_path)[1].lower()
    if extension == ".pdf":
        generated = generate_official_pdf(output_path, template.get("name") or template.get("template_type"),
                                          subject, body, metadata=metadata)
    else:
        generated = generate_official_docx(output_path, template.get("name") or template.get("template_type"),
                                           subject, body, metadata=metadata)
    stored_content = body + f"\n\nکد اعتبارسنجی: {generated['verification_token']}"
    document_id = db.save_generated_document(
        template_id=template_id, title=subject or template["name"], content=stored_content,
        file_path=os.path.abspath(output_path), zone_id=zone_id,
        related_entity_type=related_entity_type, related_entity_id=related_entity_id,
    )
    return {"id": document_id, "subject": subject, "body": body, **generated}


def generate_correspondence_letter_document(db, letter_id, output_path):
    letter = db.get_correspondence_letter(letter_id)
    if not letter:
        raise ValueError("نامه انتخاب‌شده پیدا نشد.")
    context = db.build_document_context(zone_id=letter.get("zone_id"), letter_id=letter_id)
    metadata = metadata_from_context(context, db=db, overrides={
        "number": letter.get("letter_number"),
        "date": letter.get("letter_date") or letter.get("received_date"),
        "recipient": letter.get("recipient"),
        "sender": letter.get("sender"),
        "attachment": str(letter.get("attachment_count") or 0) if letter.get("attachment_count") else "ندارد",
    })
    body = letter.get("description") or "احتراماً، به استحضار می‌رساند موضوع فوق جهت بررسی و اقدام مقتضی ارسال می‌گردد."
    extension = os.path.splitext(output_path)[1].lower()
    if extension == ".pdf":
        generated = generate_official_pdf(output_path, "نامه اداری", letter.get("subject"), body, metadata=metadata)
    else:
        generated = generate_official_docx(output_path, "نامه اداری", letter.get("subject"), body, metadata=metadata)
    document_id = db.save_generated_document(
        template_id=None, title=letter.get("subject") or "نامه اداری",
        content=body + f"\n\nکد اعتبارسنجی: {generated['verification_token']}",
        file_path=os.path.abspath(output_path), zone_id=letter.get("zone_id"),
        related_entity_type="letter", related_entity_id=letter_id,
    )
    return {"id": document_id, "subject": letter.get("subject"), "body": body, **generated}
